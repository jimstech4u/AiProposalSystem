# Builds "Thesis_Corrections.docx" — original (excerpt) + corrected text/tables
# for every section of the thesis that did not match the implemented system.
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = r"C:\Users\ajibe\Desktop\25626_final_project\Ai Proposal System\thesis\Thesis_Corrections.docx"

doc = Document()
# Base styling
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

GREY = RGBColor(0x66, 0x66, 0x66)
RED = RGBColor(0xB0, 0x00, 0x00)
GREEN = RGBColor(0x0B, 0x6A, 0x0B)


def h(text, level=1):
    doc.add_heading(text, level=level)


def para(text='', italic=False, bold=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)
    return p


def label(text, color):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = color
    return p


def original(text):
    label('Original (what the thesis currently says):', RED)
    para(text, italic=True, color=GREY)


def corrected_label():
    label('Corrected:', GREEN)


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')


def table(rows, widths=None):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Table Grid'
    for ci, cell in enumerate(t.rows[0].cells):
        cell.paragraphs[0].add_run(rows[0][ci]).bold = True
    for ri in range(1, len(rows)):
        for ci in range(len(rows[ri])):
            t.rows[ri].cells[ci].text = rows[ri][ci]
    if widths:
        for ri in t.rows:
            for ci, w in enumerate(widths):
                ri.cells[ci].width = Inches(w)
    doc.add_paragraph()
    return t


# ===================== TITLE =====================
title = doc.add_heading('Thesis Corrections — Affected Sections (Original vs Corrected)', level=0)
para('Design and Implementation of an AI-Powered Technical Proposal and Cost Estimation System — Jimstech Innovations Nigeria Limited.', italic=True, color=GREY)
para('This document contains only the sections of the thesis that did not match the implemented system. For each, it shows what the original text said and the corrected replacement text (with tables). Ground truth: roles = Software Engineer, Project Manager, Sales, System Administrator (no "Client" login); AI = NVIDIA LLM API (meta/llama-3.3-70b-instruct) with local fallbacks; database = Supabase PostgreSQL; currency = NGN (₦).')
doc.add_paragraph()

# ===================== ABSTRACT =====================
h('Abstract', 1)
original('The system was developed following an Incremental Software Development Model and incorporates machine learning-assisted estimation logic, a structured proposal generation engine, a project repository for historical reference, and a role-based user management module. The implementation leverages modern web technologies for the frontend and a robust backend supported by a relational database management system.')
corrected_label()
para('The system was developed following an Incremental Software Development Model and incorporates Large Language Model (LLM) assisted analysis, generation, and estimation — accessed through the NVIDIA NIM API (model meta/llama-3.3-70b-instruct) with deterministic local fallbacks — together with a structured proposal generation engine, a project repository for historical reference, and a role-based user management module. The implementation is a React 18 and TypeScript single-page application for the frontend, backed by Supabase (a managed PostgreSQL platform) providing authentication, row-level security policies, and audit persistence.')

# ===================== AI APPROACH (cross-cutting) =====================
h('How the AI is Implemented (LLM API, not trained ML/NLP)', 1)
para('This correction applies wherever the thesis refers to "machine-learning models trained on historical data", "NLP-based feature extraction", "classification models", or "model retraining" (Abstract, Objectives, Architecture, Testing, Recommendations).', italic=True, color=GREY)
corrected_label()
para('The intelligence in the system is provided by a Large Language Model (LLM) accessed over a hosted API, not by locally trained machine-learning models. Every AI feature — requirement analysis, proposal generation, cost estimation, timeline prediction, and technology recommendation — composes a structured prompt, sends it to the LLM, and parses the returned JSON/text.')
bullets([
    'Provider / model: NVIDIA NIM (OpenAI-compatible Chat Completions API), model meta/llama-3.3-70b-instruct, configured through VITE_NVIDIA_MODEL / VITE_NVIDIA_API_KEY.',
    'Client module: src/lib/nvidia.ts (generateWithNvidia) handles requests, authentication, errors, and retries.',
    'Resilience: when the API is unavailable, each module falls back to deterministic local heuristics so the workflow never blocks.',
    'Historical context: past projects in the repository are supplied to the model as prompt context; the model itself is not retrained. "Continuous improvement" therefore means prompt and context refinement, not model retraining.',
])

# ===================== ROLES / FUNCTIONAL REQUIREMENTS =====================
h('Functional Requirements — User Roles (REQ 1 & REQ 8)', 1)
original('The system must provide a secure login mechanism supporting multiple user roles including System Administrator, Project Manager, Software Engineer, and Client. ... Each role within the system, including System Administrator, Project Manager, Software Engineer, and Client, must be associated with a clearly defined set of permissions.')
corrected_label()
para('The system must provide a secure login mechanism supporting four user roles: System Administrator, Project Manager, Software Engineer, and Sales. Each of these roles is associated with a clearly defined set of permissions that govern which modules, records, and actions are accessible. (Note: a "client" is a data record managed by Project Manager / Sales / Admin, not a system login role.)')
para('Role permissions implemented in src/lib/permissions.ts:', bold=True)
table([
    ['Resource', 'Engineer', 'Project Manager', 'Sales', 'Admin'],
    ['Dashboard', 'read', 'read', 'read', 'read'],
    ['Requirements', 'C R U D', 'C R U D', 'C R U', 'C R U D'],
    ['Proposals', 'C R U', 'C R U D', 'C R U', 'C R U D'],
    ['Proposal Reviews', '—', 'full + approve', '—', 'full + approve'],
    ['Repository', 'read', 'C R U D', 'read', 'C R U D'],
    ['Clients', '—', 'C R U D', 'C R U D', 'C R U D'],
    ['Reports', '—', 'C R U D', 'C R U', 'C R U D'],
    ['Templates', '—', 'C R U D', '—', 'C R U D'],
    ['Integrations', '—', '—', '—', 'C R U D'],
    ['Security Audit', '—', '—', '—', 'C R U D'],
    ['Settings / Profile', 'read/update', 'read/update', 'read/update', 'full'],
])

# ===================== USE CASE DIAGRAM / TABLES =====================
h('Use Case Diagram — Actors and Use-Case Tables', 1)
original('Actors shown/described as "Client, Software Engineer, System Administrator" (Project Manager and Sales missing; Client is not a login role). Table 1 actor: "New User (Client, Software Engineer, System Administrator)"; Table 4 (Submit Requirements) actor: "Client"; Table 10 (Manage Repository) actor: "Software Engineer, System Administrator"; Table 11 (Reports) actor: "System Administrator, Software Engineer"; Tables 2 and 3 are duplicates (both "Login into System").')
corrected_label()
bullets([
    'Actors: Software Engineer, Project Manager, Sales, System Administrator, plus AI Engine (NVIDIA LLM API). Remove the "Client" actor.',
    'Table 1 (Register User Account): actor = Software Engineer / Project Manager / Sales (Admin accounts are seeded by the administrator).',
    'Table 4 (Submit Project Requirements): actor = Software Engineer / Project Manager / Sales (requirements are captured by staff, not by an external client login).',
    'Table 10 (Manage Project Repository): actor = Project Manager / System Administrator (Engineer is read-only).',
    'Table 11 (Generate Reports & Analytics): actor = Project Manager / Sales / System Administrator (not Engineer).',
    'Remove the duplicated Table 3 (identical to Table 2 "Login into System") or replace it with the intended distinct use case.',
])
para('See corrected_images/figure_07_use_case_diagram_corrected.png for the regenerated diagram.', italic=True, color=GREY)

# ===================== TECHNOLOGIES AND TOOLS USED =====================
h('Technologies and Tools Used', 1)
original('The table omits the AI engine entirely (no NVIDIA/LLM row), and does not list MUI or Framer Motion.')
corrected_label()
table([
    ['Technology / Tool', 'Use in the System'],
    ['React 18 + TypeScript', 'Frontend component architecture and typed UI implementation'],
    ['Vite', 'Development server and production build tool'],
    ['React Router DOM v7', 'Client-side routing and protected (role-based) route flow'],
    ['Tailwind CSS v4', 'Responsive styling and consistent visual design'],
    ['Radix UI', 'Accessible, unstyled UI primitives'],
    ['MUI (@mui/material)', 'Additional UI components'],
    ['Framer Motion (motion)', 'UI animation and transitions'],
    ['Lucide React', 'Application icons'],
    ['Recharts', 'Charts for dashboards, cost breakdowns, and analysis results'],
    ['Sonner', 'Toast notifications'],
    ['NVIDIA NIM LLM API (meta/llama-3.3-70b-instruct)', 'AI requirement analysis, proposal generation, cost & timeline estimation, technology recommendation'],
    ['Supabase (PostgreSQL)', 'Authentication, database, Row-Level Security policies, audit persistence'],
    ['Supabase Studio', 'Database administration and inspection'],
    ['Visual Studio Code', 'Development environment'],
], widths=[2.6, 4.0])

# ===================== SYSTEM ARCHITECTURE =====================
h('System Architectural Design', 1)
original('Presentation Layer "built using HTML, CSS, JavaScript, and Bootstrap"; "Database Layer: Built on MySQL"; "AI Processing Layer: machine learning estimation models, NLP-based requirement analysis engine"; communication via a custom PHP/REST backend.')
corrected_label()
para('The system follows a modern client + Backend-as-a-Service (BaaS) architecture composed of the following layers (this matches Figure 23):')
bullets([
    'Presentation Layer: a single-page application built with React 18 and TypeScript, bundled by Vite, styled with Tailwind CSS v4, and composed from Radix UI and MUI components. Role-specific views are rendered for Software Engineers, Project Managers, Sales, and System Administrators.',
    'Application Logic Layer: business logic runs in the browser (React hooks/services) and coordinates requirement submission, proposal generation, cost estimation, timeline prediction, and technology recommendation; role-based access is enforced in the UI (src/lib/permissions.ts).',
    'AI-Assisted Logic Layer: calls the NVIDIA LLM API (src/lib/nvidia.ts) for analysis and generation, with deterministic local fallbacks.',
    'Data & Authentication Layer: Supabase provides managed PostgreSQL, authentication (email/password with JWT sessions), and Row-Level Security (RLS) policies. The SPA talks to the database through Supabase PostgREST; database triggers maintain the audit log and notifications.',
    'External Services (future): optional CRM, e-mail, calendar, and project-management integrations.',
])
para('Key qualities: scalability (managed PostgreSQL + stateless SPA), security (RLS + JWT + HTTPS), maintainability (modular components), and reliability (local AI fallbacks and error handling).')

# ===================== DATA DICTIONARY =====================
h('Data Dictionary (Tables 16–21)', 1)
original('The data dictionary uses MySQL types (INT/AUTO_INCREMENT, TINYINT, DATETIME, short VARCHARs), includes a "password" column in the user table, sets currency DEFAULT \'USD\', and uses field names that differ from the actual schema. The real database is PostgreSQL (Supabase) with uuid keys, timestamptz, boolean, numeric, jsonb, and enums; passwords are held by Supabase Auth.')

para('Enumerated types:', bold=True)
bullets([
    'app_role: engineer, project-manager, sales, admin',
    'project_status: draft, analysis, proposal, approved, won, lost, archived',
    'proposal_status: draft, in_review, approved, rejected, sent, accepted, declined',
])

para('Original Table 16 (User) — as written in the thesis (MySQL):', bold=True, color=GREY)
table([
    ['Field', 'Type', 'Constraints', 'Description'],
    ['user_id', 'INT', 'PK, Auto Increment', 'Unique identifier'],
    ['full_name', 'VARCHAR(30)', 'NOT NULL', 'Full name'],
    ['email', 'VARCHAR(30)', 'UNIQUE, NOT NULL', 'Login email'],
    ['password', 'VARCHAR(55)', 'NOT NULL', 'Hashed password'],
    ['role', 'VARCHAR(40)', 'NOT NULL', 'User role'],
    ['is_active', 'TINYINT(1)', 'DEFAULT 1', 'Account status'],
    ['created_at', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', 'Created'],
])

corrected_label()
para('Table 16: user_profiles (PostgreSQL). Passwords are stored by Supabase Auth in auth.users — not here.', bold=True)
table([
    ['Field', 'Data Type', 'Constraints', 'Description'],
    ['id', 'uuid', 'PK, FK → auth.users(id)', 'User identifier (matches Supabase Auth)'],
    ['full_name', 'text', 'NOT NULL', 'Full name of the user'],
    ['email', 'text', 'UNIQUE, NOT NULL', 'Login email address'],
    ['phone', 'text', 'nullable', 'Contact phone number'],
    ['employee_id', 'text', 'UNIQUE, nullable', 'Employee identifier'],
    ['company_name', 'text', 'nullable', 'Company / organisation'],
    ['job_title', 'text', 'nullable', 'Job title'],
    ['department', 'text', 'nullable', 'Department'],
    ['role', 'app_role', 'NOT NULL, default engineer', 'engineer / project-manager / sales / admin'],
    ['avatar_url', 'text', 'nullable', 'Profile picture URL'],
    ['is_active', 'boolean', 'NOT NULL, default true', 'Account active status'],
    ['created_at', 'timestamptz', 'NOT NULL, default now()', 'Creation timestamp'],
    ['updated_at', 'timestamptz', 'NOT NULL, default now()', 'Last update timestamp'],
])

para('Table 17: projects (PostgreSQL).', bold=True)
table([
    ['Field', 'Data Type', 'Constraints', 'Description'],
    ['id', 'uuid', 'PK, default gen_random_uuid()', 'Project identifier'],
    ['client_id', 'uuid', 'FK → clients(id)', 'Associated client record'],
    ['title', 'text', 'NOT NULL', 'Project title'],
    ['description', 'text', 'nullable', 'Project description'],
    ['industry', 'text', 'nullable', 'Client industry domain'],
    ['project_type', 'text', 'NOT NULL', 'Type of project'],
    ['status', 'project_status', 'NOT NULL, default draft', 'Project lifecycle status'],
    ['requirements_text', 'text', 'nullable', 'Raw requirements text'],
    ['target_users', 'text', 'nullable', 'Expected user base'],
    ['integration_needs', 'text', 'nullable', 'Integration requirements'],
    ['constraints', 'text', 'nullable', 'Project constraints'],
    ['complexity_score', 'numeric(5,2)', 'CHECK 0–100', 'AI-computed complexity score'],
    ['confidence_score', 'numeric(5,2)', 'CHECK 0–100', 'Analysis confidence score'],
    ['submitted_by', 'uuid', 'FK → user_profiles(id)', 'Submitting user'],
    ['created_at / updated_at', 'timestamptz', 'default now()', 'Timestamps'],
])

para('Table 18: proposals (PostgreSQL).', bold=True)
table([
    ['Field', 'Data Type', 'Constraints', 'Description'],
    ['id', 'uuid', 'PK', 'Proposal identifier'],
    ['project_id', 'uuid', 'NOT NULL, FK → projects(id)', 'Associated project'],
    ['created_by', 'uuid', 'FK → user_profiles(id)', 'Author (was generated_by)'],
    ['title', 'text', 'NOT NULL', 'Proposal title'],
    ['template_name', 'text', 'nullable', 'Template used'],
    ['tone', 'text', "default 'professional'", 'Writing tone'],
    ['detail_level', 'text', "default 'standard'", 'Level of detail'],
    ['executive_summary / technical_approach / architecture_description', 'text', 'nullable', 'Narrative sections'],
    ['deliverables / assumptions / acceptance_criteria', 'jsonb', "default '[]'", 'Structured list sections'],
    ['status', 'proposal_status', 'NOT NULL, default draft', 'Proposal status'],
    ['version', 'integer', 'NOT NULL, default 1', 'Version number'],
    ['generated_content', 'jsonb', "default '{}'", 'Full generated content'],
    ['created_at / updated_at', 'timestamptz', 'default now()', 'Timestamps'],
])

para('Table 19: cost_estimations (+ cost_estimation_items) (PostgreSQL). Note currency default is NGN, not USD.', bold=True)
table([
    ['Field', 'Data Type', 'Constraints', 'Description'],
    ['id', 'uuid', 'PK', 'Estimation identifier'],
    ['project_id', 'uuid', 'NOT NULL, FK → projects(id)', 'Associated project'],
    ['proposal_id', 'uuid', 'FK → proposals(id)', 'Associated proposal (optional)'],
    ['created_by', 'uuid', 'FK → user_profiles(id)', 'Author'],
    ['currency', 'text', "NOT NULL, default 'NGN'", 'Currency of the estimate'],
    ['development_cost', 'numeric(14,2)', 'default 0', 'Development labour cost'],
    ['infrastructure_cost', 'numeric(14,2)', 'default 0', 'Infrastructure / hosting cost'],
    ['third_party_cost', 'numeric(14,2)', 'default 0', 'Third-party / licensing cost'],
    ['contingency_percent', 'numeric(5,2)', 'default 10', 'Contingency percentage'],
    ['contingency_amount', 'numeric(14,2)', 'default 0', 'Contingency amount'],
    ['total_cost', 'numeric(14,2)', 'default 0', 'Total projected cost'],
    ['min_cost / max_cost', 'numeric(14,2)', 'nullable', 'Estimate range'],
    ['confidence_score', 'numeric(5,2)', 'CHECK 0–100', 'Estimation confidence'],
    ['assumptions', 'jsonb', "default '[]'", 'Estimation assumptions'],
    ['created_at', 'timestamptz', 'default now()', 'Timestamp'],
])
para('cost_estimation_items: id (uuid PK), estimation_id (uuid FK), module_name (text), resource_role (text), hours (numeric), hourly_rate (numeric), multiplier (numeric), amount (numeric, generated = hours × hourly_rate × multiplier).')

para('Table 20: tech_recommendations (PostgreSQL).', bold=True)
table([
    ['Field', 'Data Type', 'Constraints', 'Description'],
    ['id', 'uuid', 'PK', 'Recommendation identifier'],
    ['project_id', 'uuid', 'NOT NULL, FK → projects(id)', 'Associated project'],
    ['created_by', 'uuid', 'FK → user_profiles(id)', 'Author'],
    ['stack_name', 'text', 'NOT NULL', 'Recommended stack name'],
    ['frontend', 'text', 'nullable', 'Recommended frontend'],
    ['backend', 'text', 'nullable', 'Recommended backend'],
    ['database_name', 'text', 'nullable', 'Recommended database'],
    ['hosting', 'text', 'nullable', 'Recommended hosting/deployment'],
    ['match_score', 'numeric(5,2)', 'CHECK 0–100', 'Suitability score'],
    ['rationale', 'text', 'nullable', 'Reasoning'],
    ['pros / cons / alternatives', 'jsonb', "default '[]'", 'Structured lists'],
    ['created_at', 'timestamptz', 'default now()', 'Timestamp'],
])

para('Table 21: notifications (PostgreSQL).', bold=True)
table([
    ['Field', 'Data Type', 'Constraints', 'Description'],
    ['id', 'uuid', 'PK', 'Notification identifier'],
    ['user_id', 'uuid', 'NOT NULL, FK → user_profiles(id)', 'Recipient'],
    ['title', 'text', 'NOT NULL', 'Notification title'],
    ['message', 'text', 'NOT NULL', 'Notification body'],
    ['type', 'text', "NOT NULL, default 'info'", 'Type (info/success/warning)'],
    ['entity_type', 'text', 'nullable', 'Related entity type'],
    ['entity_id', 'uuid', 'nullable', 'Related entity id'],
    ['read_at', 'timestamptz', 'nullable', 'Read timestamp (null = unread)'],
    ['created_at', 'timestamptz', 'default now()', 'Created timestamp'],
])
para('The full schema contains 19 tables, including clients, requirements, proposal_versions, proposal_reviews, timeline_predictions, timeline_phases, project_repository, proposal_templates, integrations, report_configs, user_settings, and audit_logs, all with RLS enabled.', italic=True, color=GREY)

# ===================== SOFTWARE TESTING =====================
h('Software Testing', 1)
original('Integration Testing: "The interface between the PHP application layer and the Python AI processing components was tested..." — there is no PHP and no local Python AI runtime.')
corrected_label()
para('Unit Testing.', bold=True)
para('Individual functions and components were tested in isolation: requirement-form input validation; the JSON-extraction helper (extractJsonObject) that parses LLM responses; the local fallback generators (cost, timeline, technology, proposal); formatCurrency (NGN formatting); and the role-permission checks in permissions.ts.')
para('Integration Testing.', bold=True)
bullets([
    'React SPA ↔ Supabase — authenticated CRUD against PostgreSQL through PostgREST, and verification that Row-Level Security correctly allows/denies operations per role.',
    'React SPA ↔ NVIDIA LLM API — request construction, response/JSON parsing, error handling (401/403/404/410/429/503), retry/back-off, and automatic switch to local fallbacks when the API is unavailable.',
    'End-to-end workflow — requirement analysis → proposal generation → cost estimation → timeline prediction → technology recommendation, verifying data flow and persistence in Supabase.',
])
para('Validation Testing.', bold=True)
para('The assembled system was validated against the functional and non-functional requirements (role-based access, NGN cost output, proposal section generation, audit logging via database triggers), confirming readiness for deployment. (Remove all references to PHP and a local Python AI runtime.)')

# ===================== HARDWARE & SOFTWARE REQUIREMENTS =====================
h('Hardware and Software Requirements', 1)
original('Server-Side: Apache 2.4 / NGINX; MySQL 8.0; PHP 8.0 with PDO; Python 3.8 with scikit-learn/NLTK/NumPy/Pandas; phpMyAdmin; an NVIDIA GPU for local model training/inference.')
corrected_label()
para('Client-Side (unchanged, still valid):', bold=True)
para('Modern browser (Chrome 90+, Firefox 88+, Edge 90+, Safari 14+); Windows 10+/macOS 11+/Linux Ubuntu 20.04+; 4 GB RAM (8 GB recommended); dual-core 2.0 GHz+; broadband ≥ 2 Mbps; 1280×720+ display.')
para('Build / Deployment (replaces the Apache/MySQL/PHP/Python/phpMyAdmin section):', bold=True)
bullets([
    'Node.js 18+ — required only to build the SPA (npm run build via Vite). No application server runs in production.',
    'Static hosting / CDN (e.g., Netlify, Vercel, Supabase Hosting, or any static web server) to serve the built dist/ assets over HTTPS.',
    'Supabase project — managed PostgreSQL 15+, Authentication, and Row-Level Security (no self-managed MySQL/Apache/PHP).',
    'Outbound HTTPS access to the NVIDIA LLM API for AI features (a valid VITE_NVIDIA_API_KEY).',
    'Database administration via Supabase Studio (not phpMyAdmin).',
    'No local GPU is required — model inference runs on NVIDIA’s hosted infrastructure.',
])

# ===================== SEQUENCE DIAGRAM DESCRIPTIONS =====================
h('Sequence-Diagram Descriptions (corrected terminology)', 1)
original('Descriptions refer to a generic "authentication module" querying a "user database", and an "NLP/ML engine". They should reference Supabase Auth and the NVIDIA LLM API.')
corrected_label()
bullets([
    'Sign Up (Fig 14): User submits the form → SPA calls Supabase Auth signUp → Supabase creates the auth user and a user_profiles row (handle_new_user trigger) → user proceeds to sign-in (OTP verification where enabled).',
    'Sign In (Fig 15): User submits email + password → SPA calls Supabase Auth signInWithPassword → Supabase returns a JWT session → SPA reads role from user_profiles → routes to the role-based dashboard.',
    'Submit & Analyze Requirements (Fig 16): Engineer/PM/Sales fills the form → SPA validates & stores in Supabase → calls the NVIDIA LLM API, which returns a structured JSON analysis (complexity, categories); local fallback on failure.',
    'Generate Proposal (Fig 17): From a selected analysis, SPA sends prompt + context to the NVIDIA LLM API → returns section content → stored in proposals / proposal_versions.',
    'Cost Estimation (Fig 18): SPA sends complexity/scope context to the NVIDIA LLM API requesting an NGN itemised estimate (JSON) → computed/normalised client-side → stored in cost_estimations + cost_estimation_items.',
    'Timeline Prediction (Fig 19): SPA requests a phased schedule from the NVIDIA LLM API → stored in timeline_predictions / timeline_phases.',
    'Send Notification (Fig 20): notifications are system-generated — database triggers (e.g., notify_proposal_review) insert rows into notifications, shown on the in-app bell. E-mail dispatch and a manual compose screen are future work.',
])

# ===================== RECOMMENDATIONS =====================
h('Recommendations — Continuous Improvement', 1)
original('"As the system accumulates data... the AI estimation models should be periodically retrained and validated against actual project outcomes... A quarterly model review process is recommended."')
corrected_label()
para('Because the system uses a hosted Large Language Model (not a locally trained model), continuous improvement is achieved through prompt and context refinement rather than retraining. As the project repository accumulates completed projects, that historical data should be curated and supplied to the model as richer prompt context, and the prompts/parameters periodically reviewed (e.g., quarterly) and validated against actual project outcomes to improve estimate accuracy. The deterministic local fallbacks should also be tuned as more real cost/timeline data becomes available.')

doc.save(OUT)
print('SAVED', OUT, os.path.getsize(OUT), 'bytes')
